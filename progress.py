import asyncio
import uuid
import time
import argparse
import re
from datetime import datetime
from playwright.async_api import async_playwright
from docx import Document
from bs4 import BeautifulSoup

SELECTORS = {
    "modal": "div._modal_7bdw1_1",
    "close_btn": "div._close_7bdw1_31",
    "prompt_textarea": "textarea",
    "submit_btn": "button[type='submit'][class*='_buttons-send-wrapper_']",
    "content_block": [
        "div._content_1k32x_12",
        "div._container_q86iu_1",
        "div[data-testid='virtuoso-item-list'] > div",
        "div._content_6r4i1_29 div._container_q86iu_1"
    ],
    "reference_scroller": "div._virtuoso_6r4i1_26",
    "reference_block": [
        "div[data-index] div._container_q86iu_1",
        "div[data-item-index] div._container_q86iu_1",
        "div._container_q86iu_1",
        "div[data-testid='virtuoso-item-list'] > div",
        "div._content_6r4i1_29 div._container_q86iu_1"
    ],
    "reference_index": "div._index_q86iu_12",
    "reference_title": "div._title-paragraph_1doxh_4 p",
    "reference_author": "div._author_name_1fn6n_38",
    "reference_journal": "span._name_niu8h_11",
    "reference_date": "div._journal-date_q86iu_51",
    "loading_spinner": None,
}

def extract_cited_reference_numbers(html_content):
    numbers = re.findall(r'\[(\d+)\]', html_content)
    return sorted(set(int(n) for n in numbers))

def is_visible_bs4(element):
    style = element.get('style', '')
    return 'display: none' not in style and 'visibility: hidden' not in style

def extract_paragraph(child):
    p_text = child.get_text(strip=True)
    return p_text if p_text else None

def extract_list(child):
    items = []
    for li in child.find_all("li", recursive=False):
        li_text = li.get_text(strip=True)
        if li_text:
            items.append(li_text)
    return items

def extract_table(child):
    rows = child.find_all("tr")
    table_data = []
    for row in rows:
        cells = row.find_all(["td", "th"])
        table_data.append([cell.get_text(strip=True) for cell in cells])
    return table_data

def extract_image_info(img_container):
    """Extracts image information from container with proper handling of nested structures."""
    images = []
    # Handle direct img tags
    for img in img_container.find_all("img", recursive=False):
        img_url = img.get("src", "")
        if not img_url:
            continue
        caption = ""
        caption_div = img_container.find("div", class_="_img-title_1k32x_79")
        if caption_div:
            caption = caption_div.get_text(strip=True)
        if not caption:
            caption = img.get("alt", "")
        source = ""
        em_tag = img_container.find_next_sibling("em")
        if not em_tag and img_container.parent:
            em_tag = img_container.parent.find("em")
        if em_tag:
            source = em_tag.get_text(strip=True)
        images.append((img_url, caption, source))
    # Handle nested image containers
    for nested in img_container.find_all("div", class_="_img_1k32x_74"):
        img = nested.find("img")
        if not img:
            continue
        img_url = img.get("src", "")
        if not img_url:
            continue
        caption = ""
        caption_div = nested.find("div", class_="_img-title_1k32x_79")
        if caption_div:
            caption = caption_div.get_text(strip=True)
        if not caption:
            caption = img.get("alt", "")
        source = ""
        em_tag = nested.find_next_sibling("em")
        if not em_tag and nested.parent:
            em_tag = nested.parent.find("em")
        if em_tag:
            source = em_tag.get_text(strip=True)
        images.append((img_url, caption, source))
    return images

def parse_and_save_content(html_content, prompt_text, references_dict, cited_numbers):
    print("[*] Parsing final content...")
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        doc.add_heading(prompt_text, level=1)
        content_blocks = []
        for selector in SELECTORS["content_block"]:
            content_blocks.extend(soup.select(selector))
        print(f"[*] Found {len(content_blocks)} content blocks in HTML.")
        inserted_content = set()
        for block_idx, block in enumerate(content_blocks):
            print(f"[*] Processing content block {block_idx+1}/{len(content_blocks)}")
            try:
                children = list(block.children)
                i = 0
                while i < len(children):
                    elem = children[i]
                    if isinstance(elem, str) or not hasattr(elem, "name"):
                        i += 1
                        continue
                    if not is_visible_bs4(elem):
                        i += 1
                        continue
                    # Paragraphs
                    if elem.name == "p":
                        p_text = extract_paragraph(elem)
                        if p_text and p_text not in inserted_content:
                            doc.add_paragraph(p_text)
                            inserted_content.add(p_text)
                            print(f"    [+] Inserted paragraph: {p_text[:60]}...")
                        # Images inside paragraph
                        img_divs = elem.find_all("div", class_="_img_1k32x_74")
                        for img_div in img_divs:
                            images = extract_image_info(img_div)
                            for img_url, caption, source in images:
                                if img_url and img_url not in inserted_content:
                                    image_info = f"Image URL: {img_url}"
                                    if caption:
                                        image_info += f"\nCaption: {caption}"
                                    if source:
                                        image_info += f"\nSource: {source}"
                                    doc.add_paragraph(image_info)
                                    inserted_content.add(img_url)
                                    print(f"    [+] Inserted image: {img_url}")
                        # Next sibling image container
                        if i + 1 < len(children):
                            next_elem = children[i + 1]
                            if hasattr(next_elem, "name") and next_elem.name == "div" and "_img_1k32x_74" in next_elem.get("class", []):
                                images = extract_image_info(next_elem)
                                for img_url, caption, source in images:
                                    if img_url and img_url not in inserted_content:
                                        image_info = f"Image URL: {img_url}"
                                        if caption:
                                            image_info += f"\nCaption: {caption}"
                                        if source:
                                            image_info += f"\nSource: {source}"
                                        doc.add_paragraph(image_info)
                                        inserted_content.add(img_url)
                                        print(f"    [+] Inserted image: {img_url}")
                                i += 1  # Skip the image container since we've processed it
                    # Lists
                    elif elem.name in ["ul", "ol"]:
                        items = extract_list(elem)
                        for li_text in items:
                            if li_text and li_text not in inserted_content:
                                doc.add_paragraph(f"- {li_text}")
                                inserted_content.add(li_text)
                                print(f"    [+] Inserted list item: {li_text[:60]}...")
                    # Tables
                    elif elem.name == "table":
                        table_data = extract_table(elem)
                        if table_data:
                            t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            for row_idx, row in enumerate(table_data):
                                for col_idx, cell_text in enumerate(row):
                                    t.cell(row_idx, col_idx).text = cell_text
                            print("    [+] Inserted table.")
                    # Standalone image containers
                    elif elem.name == "div" and "_img_1k32x_74" in elem.get("class", []):
                        images = extract_image_info(elem)
                        for img_url, caption, source in images:
                            if img_url and img_url not in inserted_content:
                                image_info = f"Image URL: {img_url}"
                                if caption:
                                    image_info += f"\nCaption: {caption}"
                                if source:
                                    image_info += f"\nSource: {source}"
                                doc.add_paragraph(image_info)
                                inserted_content.add(img_url)
                                print(f"    [+] Inserted image: {img_url}")
                    i += 1
            except Exception as e:
                print(f"    [ERROR] Failed to process block: {e}")
        # Add references section
        if references_dict:
            doc.add_heading("References", level=2)
            for ref_num in cited_numbers:
                ref_str = str(ref_num)
                if ref_str in references_dict:
                    doc.add_paragraph(f"[{ref_str}] {references_dict[ref_str]}")
            print(f"[*] Added {len(cited_numbers)} cited references to document.")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = uuid.uuid4().hex[:6]
        filename = f"bohrium_ai_response_{timestamp}_{unique_id}.docx"
        try:
            doc.save(filename)
            print(f"[✓] Document saved as '{filename}'")
        except Exception as e:
            print(f"[ERROR] Failed to save document: {e}")
    except Exception as e:
        print(f"[ERROR] Exception during parsing and saving content: {e}")

async def close_modal(page):
    try:
        modal = await page.query_selector(SELECTORS["modal"])
        if modal:
            print("[!] Login popup detected.")
            close_btn = await modal.query_selector(SELECTORS["close_btn"])
            if close_btn:
                await close_btn.click()
                print("[!] Login popup closed.")
                await page.wait_for_selector(SELECTORS["modal"], state="detached", timeout=7000)
    except Exception as e:
        print(f"[WARNING] Could not close modal: {e}")

async def enter_prompt(page, prompt_text):
    try:
        print("[*] Waiting for prompt textarea...")
        await page.wait_for_selector(SELECTORS["prompt_textarea"], timeout=15000)
        await page.fill(SELECTORS["prompt_textarea"], prompt_text)
        print(f"[*] Prompt entered: '{prompt_text}'")
        await page.wait_for_selector(SELECTORS["submit_btn"], timeout=15000)
        submit_btn = await page.query_selector(SELECTORS["submit_btn"])
        await submit_btn.scroll_into_view_if_needed()
        print("[*] Clicking submit button...")
        await submit_btn.click()
        print("[*] Query submitted.")
    except Exception as e:
        print(f"[ERROR] Failed to enter prompt: {e}")
        raise

async def wait_for_content(page):
    print("[*] Waiting 20s before content processing...")
    await asyncio.sleep(20)
    stable_count = 0
    last_len = 0
    stable_threshold = 10
    max_wait = 600
    pre_refresh_wait = 10
    post_refresh_wait = 10
    refreshed = False
    start_time = time.time()
    while True:
        try:
            elapsed = time.time() - start_time
            if elapsed > max_wait:
                print("[!] Max wait reached.")
                break
            if int(elapsed) % 10 == 0:
                try:
                    await page.evaluate("window.scrollBy(0, window.innerHeight);")
                    print("[*] Scrolled page to load more content.")
                except Exception as e:
                    print(f"[WARNING] Failed to scroll: {e}")
                await asyncio.sleep(1)
            elements = []
            for selector in SELECTORS["content_block"]:
                try:
                    elements.extend(await page.query_selector_all(selector))
                except Exception as e:
                    continue
            print(f"[ ] {int(elapsed)}s: Found {len(elements)} content blocks.")
            try:
                combined_html = "".join([await el.evaluate("node => node.outerHTML") for el in elements])
                print(f"[ ] Combined HTML length: {len(combined_html)}")
            except Exception as e:
                print(f"[ERROR] Failed to combine HTML: {e}")
                combined_html = ""
            if len(combined_html) == last_len:
                stable_count += 1
                print(f"[✓] Stable for {stable_count}s.")
            else:
                stable_count = 0
                print("[*] Content changed, resetting stability counter.")
            last_len = len(combined_html)
            if stable_count >= stable_threshold and not refreshed:
                print(f"[✓] Stability detected. Waiting {pre_refresh_wait}s before refreshing...")
                await asyncio.sleep(pre_refresh_wait)
                print("[✓] Performing one-time refresh now...")
                try:
                    await page.reload()
                    print("[*] Page reloaded.")
                except Exception as e:
                    print(f"[ERROR] Failed to reload page: {e}")
                if SELECTORS["loading_spinner"]:
                    try:
                        await page.wait_for_selector(SELECTORS["loading_spinner"], state="detached", timeout=10000)
                    except Exception as e:
                        print(f"[WARNING] Loading spinner not found after refresh: {e}")
                print(f"[*] Waiting {post_refresh_wait}s after refresh...")
                await asyncio.sleep(post_refresh_wait)
                refreshed = True
                break
            await asyncio.sleep(1)
        except Exception as e:
            print(f"[ERROR] An error occurred while waiting for content: {e}")
            break

async def extract_cited_references(page, cited_numbers):
    print("[*] Extracting cited references...")
    references_dict = {}
    found_refs = set()
    max_scroll_attempts = max(cited_numbers) * 3 if cited_numbers else 20
    scroll_attempts = 0
    consecutive_no_new = 0
    max_consecutive_no_new = 10
    scroll_step = 400

    ref_scroller = await page.query_selector(SELECTORS["reference_scroller"])
    while set(cited_numbers).difference(found_refs) and scroll_attempts < max_scroll_attempts and consecutive_no_new < max_consecutive_no_new:
        sources_before = len(found_refs)
        for selector in SELECTORS["reference_block"]:
            ref_blocks = await page.query_selector_all(selector)
            for ref_block in ref_blocks:
                try:
                    ref_num_tag = await ref_block.query_selector(SELECTORS["reference_index"])
                    ref_num = await ref_num_tag.inner_text() if ref_num_tag else None
                    if ref_num:
                        ref_num = ref_num.replace('.', '').strip()
                    if not ref_num or not ref_num.isdigit():
                        continue
                    ref_num_int = int(ref_num)
                    if ref_num_int not in cited_numbers or ref_num in references_dict:
                        continue
                    title_tag = await ref_block.query_selector(SELECTORS["reference_title"])
                    title = await title_tag.inner_text() if title_tag else ""
                    author_tags = await ref_block.query_selector_all(SELECTORS["reference_author"])
                    authors = ", ".join([await a.inner_text() for a in author_tags]) if author_tags else ""
                    date_tag = await ref_block.query_selector(SELECTORS["reference_date"])
                    date = await date_tag.inner_text() if date_tag else ""
                    journal_tag = await ref_block.query_selector(SELECTORS["reference_journal"])
                    journal = await journal_tag.inner_text() if journal_tag else ""
                    ref_text = f"{authors}. {date}. {title}. {journal}."
                    references_dict[ref_num] = ref_text
                    found_refs.add(ref_num_int)
                    print(f"    [+] Extracted cited reference [{ref_num}]: {ref_text}")
                except Exception as e:
                    continue
        sources_after = len(found_refs)
        if sources_after > sources_before:
            consecutive_no_new = 0
        else:
            consecutive_no_new += 1
        if set(cited_numbers).issubset(found_refs):
            print("[*] All cited references extracted.")
            break
        if ref_scroller:
            await ref_scroller.evaluate(f"el => el.scrollBy(0, {scroll_step})")
        else:
            await page.evaluate(f"window.scrollBy(0, {scroll_step});")
        await asyncio.sleep(0.7)
        scroll_attempts += 1
    return references_dict

async def run_bohrium_search(prompt_text, headless):
    async with async_playwright() as p:
        print("[*] Launching browser...")
        try:
            browser = await p.chromium.launch(headless=headless)
            context = await browser.new_context()
            page = await context.new_page()
        except Exception as e:
            print(f"[ERROR] Failed to launch browser: {e}")
            return
        try:
            print("[*] Navigating to Bohrium AI...")
            await page.goto("https://www.bohrium.com/en-US", timeout=60000, wait_until="domcontentloaded")
            await page.wait_for_timeout(6000)
            close_modal_task = asyncio.create_task(close_modal(page))
            print("[*] Page loaded.")
            await enter_prompt(page, prompt_text)
            await wait_for_content(page)
            print("[*] Collecting main content for reference scan...")
            try:
                elements = []
                for selector in SELECTORS["content_block"]:
                    elements.extend(await page.query_selector_all(selector))
                combined_html_after_refresh = "".join([await el.evaluate("node => node.outerHTML") for el in elements])
            except Exception as e:
                print(f"[ERROR] Failed to extract refreshed content: {e}")
                combined_html_after_refresh = ""
            cited_numbers = extract_cited_reference_numbers(combined_html_after_refresh)
            print(f"[*] Cited reference numbers in content: {cited_numbers}")
            references_dict = await extract_cited_references(page, cited_numbers)
            parse_and_save_content(combined_html_after_refresh, prompt_text, references_dict, cited_numbers)
            await page.wait_for_timeout(3000)
        except Exception as e:
            print(f"[ERROR] Exception occurred: {e}")
        finally:
            print("[*] Closing browser...")
            try:
                await browser.close()
            except Exception as e:
                print(f"[ERROR] Failed to close browser: {e}")
            if 'close_modal_task' in locals() and not close_modal_task.done():
                close_modal_task.cancel()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run a search on Bohrium AI and save the results.")
    parser.add_argument("prompt", type=str, help="The search prompt to use.")
    parser.add_argument("--headless", action="store_true", help="Run the browser in headless mode.")
    args = parser.parse_args()
    try:
        asyncio.run(run_bohrium_search(args.prompt, args.headless))
    except Exception as e:
        print(f"[ERROR] An error occurred while running the script: {e}")
