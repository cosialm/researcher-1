import asyncio
import uuid
import time
import argparse
from datetime import datetime
from playwright.async_api import async_playwright
from docx import Document
from bs4 import BeautifulSoup

# Configuration for CSS selectors
SELECTORS = {
    "modal": "div._modal_7bdw1_1",
    "close_btn": "div._close_7bdw1_31",
    "prompt_textarea": "textarea",
    "submit_btn": "button._buttons-send-wrapper_bf84e_11",
    "content_block": "div.mt-20",
    "image_div": "div._img_1k32x_74",
    "image_tag": "img",
    "caption_div": "div._img-title_1k32x_79",
    "source_tag": "em",
    "loading_spinner": None,  # optional: e.g. "div.loading-spinner"
}


async def close_modal(page):
    """Closes the login popup if it appears."""
    try:
        modal = await page.query_selector(SELECTORS["modal"])
        if modal:
            print("[!] Login popup detected.")
            close_btn = await modal.query_selector(SELECTORS["close_btn"])
            if close_btn:
                await close_btn.click()
                print("[!] Login popup closed.")
                await page.wait_for_selector(SELECTORS["modal"], state="detached", timeout=7000)
    except Exception as e:
        print(f"[WARNING] Could not close modal: {e}")


async def enter_prompt(page, prompt_text):
    """Enters the prompt into the textarea and submits."""
    try:
        await page.wait_for_selector(SELECTORS["prompt_textarea"], timeout=15000)
        await page.fill(SELECTORS["prompt_textarea"], prompt_text)
        print(f"[*] Prompt entered: '{prompt_text}'")
        await page.wait_for_selector(SELECTORS["submit_btn"], timeout=15000)
        submit_btn = await page.query_selector(SELECTORS["submit_btn"])
        await submit_btn.scroll_into_view_if_needed()
        print("[*] Clicking submit button...")
        await submit_btn.click()
        print("[*] Query submitted.")
    except Exception as e:
        print(f"[ERROR] Failed to enter prompt: {e}")
        raise


async def wait_for_content(page):
    """Waits for the content to load and stabilize."""
    print("[*] Waiting 20s before content processing...")
    await asyncio.sleep(20)

    stable_count = 0
    last_len = 0
    stable_threshold = 10
    max_wait = 600
    pre_refresh_wait = 10
    post_refresh_wait = 10
    refreshed = False
    start_time = time.time()

    while True:
        try:
            elapsed = time.time() - start_time
            if elapsed > max_wait:
                print("[!] Max wait reached.")
                break

            if int(elapsed) % 10 == 0:
                await page.evaluate("window.scrollBy(0, window.innerHeight);")
                await asyncio.sleep(1)

            elements = await page.query_selector_all(SELECTORS["content_block"])
            print(f"[ ] {int(elapsed)}s: Found {len(elements)} content blocks.")

            combined_html = "".join([await el.inner_html() for el in elements])

            if len(combined_html) == last_len:
                stable_count += 1
                print(f"[✓] Stable for {stable_count}s.")
            else:
                stable_count = 0
                print("[*] Content changed, resetting stability counter.")

            last_len = len(combined_html)

            if stable_count >= stable_threshold and not refreshed:
                print(f"[✓] Stability detected. Waiting {pre_refresh_wait}s before refreshing...")
                await asyncio.sleep(pre_refresh_wait)
                print("[✓] Performing one-time refresh now...")
                await page.reload()
                if SELECTORS["loading_spinner"]:
                    try:
                        await page.wait_for_selector(SELECTORS["loading_spinner"], state="detached", timeout=10000)
                    except Exception as e:
                        print(f"[WARNING] Loading spinner not found after refresh: {e}")
                print(f"[*] Waiting {post_refresh_wait}s after refresh...")
                await asyncio.sleep(post_refresh_wait)
                refreshed = True
                break

            await asyncio.sleep(1)
        except Exception as e:
            print(f"[ERROR] An error occurred while waiting for content: {e}")
            break


def parse_and_save_content(html_content, prompt_text):
    """Parses the HTML content and saves it to a .docx file."""
    print("[*] Parsing final content...")
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    doc.add_heading(prompt_text, level=1)
    captions_already_inserted = set()

    for block in soup.select(SELECTORS["content_block"]):
        paragraphs = block.find_all("p", recursive=True)
        if not paragraphs:
            text = block.get_text(strip=True)
            if text and text not in captions_already_inserted:
                doc.add_paragraph(text)
                print(f"[+] Inserted block text: {text[:60]}...")
            continue

        for p in paragraphs:
            img_div = p.find("div", class_=SELECTORS["image_div"].split('.')[-1])
            if img_div:
                img_tag = img_div.find(SELECTORS["image_tag"])
                img_url = ""
                if img_tag and img_tag.has_attr("src"):
                    img_url = img_tag["src"]

                caption_div = img_div.find("div", class_=SELECTORS["caption_div"].split('.')[-1])
                caption_text = ""
                if caption_div:
                    caption_text = caption_div.get_text(strip=True)

                source_text = ""
                em_tag = p.find(SELECTORS["source_tag"])
                if em_tag:
                    source_text = em_tag.get_text(strip=True)

                full_sentence = f"{img_url}. {caption_text}. Source: {source_text}."
                if full_sentence not in captions_already_inserted:
                    doc.add_paragraph(full_sentence)
                    captions_already_inserted.add(full_sentence)
                    print(f"[+] Inserted image+caption+source: {full_sentence[:60]}...")
                continue

            p_text = p.get_text(strip=True)
            if p_text and p_text not in captions_already_inserted:
                doc.add_paragraph(p_text)
                captions_already_inserted.add(p_text)
                print(f"[+] Inserted paragraph: {p_text[:60]}...")
            elif p_text:
                print(f"[–] Skipped duplicate: {p_text[:60]}...")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex[:6]
    filename = f"bohrium_ai_response_{timestamp}_{unique_id}.docx"
    doc.save(filename)
    print(f"[✓] Document saved as '{filename}'")


async def run_bohrium_search(prompt_text, headless):
    """Main function to run the Bohrium search."""
    async with async_playwright() as p:
        print("[*] Launching browser...")
        browser = await p.chromium.launch(headless=headless)
        context = await browser.new_context()
        page = await context.new_page()

        try:
            print("[*] Navigating to Bohrium AI...")
            await page.goto("https://www.bohrium.com/en-US", timeout=60000, wait_until="domcontentloaded")
            await page.wait_for_timeout(6000)

            # Asynchronously close modals
            close_modal_task = asyncio.create_task(close_modal(page))

            print("[*] Page loaded.")

            await enter_prompt(page, prompt_text)
            await wait_for_content(page)

            print("[*] Collecting final refreshed content...")
            refreshed_elements = await page.query_selector_all(SELECTORS["content_block"])
            combined_html_after_refresh = "".join([await el.inner_html() for el in refreshed_elements])

            parse_and_save_content(combined_html_after_refresh, prompt_text)

            await page.wait_for_timeout(3000)

        except Exception as e:
            print(f"[ERROR] Exception occurred: {e}")
        finally:
            print("[*] Closing browser...")
            await browser.close()
            if 'close_modal_task' in locals() and not close_modal_task.done():
                close_modal_task.cancel()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run a search on Bohrium AI and save the results.")
    parser.add_argument("prompt", type=str, help="The search prompt to use.")
    parser.add_argument("--headless", action="store_true", help="Run the browser in headless mode.")
    args = parser.parse_args()

    asyncio.run(run_bohrium_search(args.prompt, args.headless))
